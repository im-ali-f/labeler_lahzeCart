import docx
#create a doc file to save used card numbers
usedDocument = docx.Document("used.docx") #ino bayad hatman akarin noskasho dashte bashi
all_usedparas=usedDocument.paragraphs
all_usedparas_array=[]
for para in all_usedparas:
    all_usedparas_array.append(para.text)

doc = docx.Document("1000 ta adad.docx") # inja file mored nazaro mizari
#doc = docx.Document("text.docx") # in baraye test
all_paras = doc.paragraphs
all_paras_array=[]
for para in all_paras:
    all_paras_array.append(para.text)

# shomare cartaii ke moshabeh nistano dar miare
def checkExists(paraToCheck):
    for usedPara in all_usedparas_array:
        if usedPara == paraToCheck:
            return
    return paraToCheck

remainParas = filter(checkExists,all_paras_array)


#for usedPara in all_usedparas:
    #(usedPara.text != para.text):


# hala sakht image
import PIL
from PIL import Image,ImageDraw,ImageFont
PIL.Image.MAX_IMAGE_PIXELS = 933120000
imgCounter=0
for para in remainParas:
    im = PIL.Image.new(mode="RGB", size=(531,118), color = "white")
    imPath=f"newCards/img{para}.jpg"
    im.save(imPath)

    # get a font
    fnt = ImageFont.truetype("Handmade.otf", 70)
    # get a drawing context
    d = ImageDraw.Draw(im)
    innerCounter=0
    cardNum=""
    for num in para:
        if innerCounter <4:
            cardNum+=num
            innerCounter+=1
        else:
            innerCounter=1
            cardNum += "-"
            cardNum+=num
    # draw multiline text
    W=531
    H=118
    _, _, w, h = d.textbbox((0, 0), cardNum, font=fnt)
    d.text(((W - w) / 2, (H - h) / 2), cardNum, font=fnt, fill=(0,0,0))
    im.save(imPath)
    #on shomare estefade shode bayad save she to used
    usedDocument.add_paragraph(para)


# inja be 70cm x 106cm tabdil mikonim

import os


directory = 'newCards'
files = os.listdir(directory)
indexFiles = 0
width=24
height=24
name=1
imA4 = PIL.Image.new(mode="RGB", size=(12508,8260), color = "black")
imA4Path = f"a4Cards/{str(name)}.png"
imA4.save(imA4Path)

heightFull=0
while indexFiles < len(files):
    if heightFull == 1:
        imA4 = PIL.Image.new(mode="RGB", size=(12508,8260), color="black")
        imA4Path = f"a4Cards/{str(name)}.png"
        imA4.save(imA4Path)
        heightFull=0
    filename = files[indexFiles]
    print(filename)
    im2=Image.open(f"newCards/{filename}")
    innerimA4=Image.open(imA4Path)
    backIm=innerimA4.copy()
    backIm.paste(im2,(width,height))
    indexFiles += 1
    if(width < 11655):
        width+=555
    else:
        width=24
        height+=142

    if (height <8094):
        imA4Path=f"a4Cards/{str(name)}.png"
        backIm.save(imA4Path)
    else:
        imA4Path=f"a4Cards/{str(name)}.png"
        backIm.save(imA4Path)
        height=24
        width=24
        name+=1
        heightFull=1
    os.remove(f"newCards/{filename}")
#akhar az hame kolan save mishe to file mon
#usedDocument.save('used.docx')